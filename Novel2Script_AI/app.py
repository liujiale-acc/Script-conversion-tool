import os
import time
import re
from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, RGBColor
from openai import OpenAI
from dotenv import load_dotenv

# 加载环境变量
load_dotenv()

app = Flask(__name__)

# 初始化 OpenAI 客户端 (DeepSeek API)
client = OpenAI(
    api_key=os.getenv("DEEPSEEK_API_KEY"),
    base_url="https://api.deepseek.com"
)

@app.route('/')
def index():
    """首页路由"""
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    """核心转换路由：支持多规格剧本生成"""
    try:
        data = request.get_json()
        raw_text = data.get('text', '')
        category = data.get('category', '电影标准模式')

        # v3.0 工业级 Prompt：集成了诊断、剧本创作、数字化映射、制片拆解
        prompt = f"""
        你是一位顶级影视制片专家、编剧和统筹副导演。
        请将以下小说内容重构为标准的【{category}】规格剧本。

        输出要求：必须严格按以下三个板块输出，板块间用标签分隔，严禁使用Markdown符号。

        [[DIAGNOSIS]]
        [剧本诊断报告]
        张力评分：(0-10分)
        节奏建议：(针对{category}规格的专业建议)
        核心冲突：(一句话总结)

        [[TEXT_VERSION]]
        [角色画像分析]
        (提取主要人物的外貌、性格特征、核心动机)

        [专业剧本正文]
        (包含场景序号、地点、时间、内/外景、人物台词及动作神态)

        [视觉分镜提示词]
        (为关键场景提供AI绘画描述词)

        [[YAML_VERSION]]
        ---
        metadata:
          title: "名称"
          category: "{category}"
        scenes:
          - scene_no: 1
            setting: {{ loc: "地点", time: "时间", type: "内景/外景" }}
            props: ["关键道具1", "重要物件2"]
            characters: ["角色A", "角色B"]
            content:
              - type: "action"
                text: "画面动作描写"
              - type: "dialogue"
                role: "角色名"
                emotion: "神态"
                text: "台词内容"
        ---
        """

        # 调用 AI
        res = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": raw_text}
            ]
        )
        
        full_output = res.choices[0].message.content
        # 预清洗：去掉可能干扰解析的 Markdown 符号
        full_output = re.sub(r'[*#]', '', full_output)
        
        return jsonify({"success": True, "script": full_output})

    except Exception as e:
        print(f"Error during conversion: {e}")
        return jsonify({"success": False, "msg": str(e)})

@app.route('/download', methods=['POST'])
def download():
    """文件下载路由：支持 Word 和 YAML"""
    try:
        data = request.get_json()
        content = data.get('content', '')
        file_type = data.get('type', 'docx')
        
        if not content:
            return jsonify({"success": False, "msg": "内容为空"}), 400

        if file_type == 'yaml':
            # 导出 YAML
            filename = f"production_data_{int(time.time())}.yaml"
            path = os.path.join('scripts', filename)
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
        else:
            # 导出 Word (保留专业格式排版)
            doc = Document()
            # 设置默认字体（针对中文）
            doc.styles['Normal'].font.name = 'SimSun'
            
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue
                p = doc.add_paragraph()
                # 针对标题行加粗处理
                if "[" in line and "]" in line:
                    run = p.add_run(line)
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(44, 62, 80)
                else:
                    p.add_run(line)
            
            filename = f"AI_Script_v3_Final_{int(time.time())}.docx"
            path = os.path.join('scripts', filename)
            doc.save(path)

        return jsonify({"success": True, "filename": filename})
    except Exception as e:
        return jsonify({"success": False, "msg": str(e)})

@app.route('/get_file/<filename>')
def get_file(filename):
    """发送文件给浏览器"""
    return send_file(os.path.join('scripts', filename), as_attachment=True)

if __name__ == '__main__':
    # 确保存储目录存在
    if not os.path.exists('scripts'):
        os.makedirs('scripts')
    app.run(debug=True, port=5000)